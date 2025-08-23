"""
RAG (Retrieval Augmented Generation) Service with PGVector
Reads configuration from database tool instances, models tables instead of environment variables.
"""

from fastapi import FastAPI, HTTPException, status, UploadFile, File, Form, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from typing import Dict, Any, List, Optional, Union
import logging
from datetime import datetime
import json
import uuid
import os
import tempfile
from contextlib import asynccontextmanager
import asyncpg
from sqlalchemy import create_engine, text
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession, async_sessionmaker
import google.generativeai as genai
import openai
import httpx
import aiofiles
from pathlib import Path
import numpy as np

# Document processing imports

import pypdf
from docx import Document as DocxDocument
# Docling integration
from docling import Document as DoclingDocument
from docling.chunking import chunk_document

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration from environment (fallback only)
DATABASE_URL = os.getenv("DATABASE_URL", "postgresql+asyncpg://agenticai_user:agenticai_password@postgres:5432/agenticai_platform")
UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "./uploads"))

# Global variables
async_engine = None
async_session_maker = None

# Configuration cache
llm_models_cache = {}
embedding_models_cache = {}
rag_tool_instances_cache = {}

async def get_database_session():
    """Dependency to get database session"""
    async with async_session_maker() as session:
        try:
            yield session
        finally:
            await session.close()

async def load_rag_configuration():
    """Load RAG configuration from database tables"""
    global llm_models_cache, embedding_models_cache, rag_tool_instances_cache
    
    try:
        async with async_session_maker() as session:
            # Load LLM models (text-generation types)
            result = await session.execute(text("""
                SELECT name, display_name, provider, endpoint_url, model_config, pricing_info
                FROM llm_models 
                WHERE is_active = true AND model_type = 'text-generation'
            """))
            llm_models_cache = {}
            for row in result:
                mconf = row[4] or {}
                capabilities = mconf.get('capabilities', {}) if isinstance(mconf, dict) else {}
                llm_models_cache[row[0]] = {
                    'name': row[0],
                    'display_name': row[1],
                    'provider': row[2],
                    'api_endpoint': row[3],
                    'capabilities': capabilities,
                    'pricing_info': row[5]
                }

            # Load embedding models
            result = await session.execute(text("""
                SELECT name, display_name, provider, endpoint_url, model_config, pricing_info
                FROM embedding_models 
                WHERE is_active = true
            """))
            embedding_models_cache = {}
            for row in result:
                mconf = row[4] or {}
                capabilities = mconf.get('capabilities', {}) if isinstance(mconf, dict) else {}
                embedding_models_cache[row[0]] = {
                    'name': row[0],
                    'display_name': row[1],
                    'provider': row[2],
                    'api_endpoint': row[3],
                    'capabilities': capabilities,
                    'pricing_info': row[5]
                }
            
            # Load RAG tool instances (could be vector store configs, API keys, etc.)
            result = await session.execute(text("""
                SELECT ti.name, ti.display_name, tt.name as template_name, ti.configuration
                FROM tool_instances ti
                JOIN tool_templates tt ON ti.tool_template_id = tt.id
                WHERE ti.status = 'active' AND (tt.name LIKE '%rag%' OR tt.name LIKE '%embedding%' OR tt.name LIKE '%vector%')
            """))
            rag_tool_instances_cache = {row[0]: {
                'name': row[0],
                'display_name': row[1],
                'template_name': row[2],
                'configuration': row[3]
            } for row in result}
            
            logger.info(f"Loaded {len(llm_models_cache)} LLM models, {len(embedding_models_cache)} embedding models, {len(rag_tool_instances_cache)} RAG instances")
            
    except Exception as e:
        logger.error(f"Error loading RAG configuration: {e}")
        # Use default configuration if database is not available
        llm_models_cache = {
            'gpt-4o': {
                'name': 'gpt-4o',
                'display_name': 'GPT-4 Omni',
                'provider': 'OpenAI',
                'api_endpoint': 'https://api.openai.com/v1/chat/completions',
                'capabilities': {'max_tokens': 128000},
                'pricing_info': {}
            }
        }
        embedding_models_cache = {
            'text-embedding-3-small': {
                'name': 'text-embedding-3-small',
                'display_name': 'OpenAI Text Embedding 3 Small',
                'provider': 'OpenAI',
                'api_endpoint': 'https://api.openai.com/v1/embeddings',
                'capabilities': {'dimensions': 1536},
                'pricing_info': {}
            }
        }
async def get_embedding_from_model(text: str, model_name: str = 'text-embedding-3-small') -> List[float]:
    """Get embedding for text using specified model"""
    if model_name not in embedding_models_cache:
        raise HTTPException(status_code=404, detail=f"Embedding model {model_name} not found")
    
    model_config = embedding_models_cache[model_name]
    
    if model_config['provider'] == 'OpenAI':
        # Check if we have API key from tool instance configuration
        api_key = None
        for instance in rag_tool_instances_cache.values():
            config = instance.get('configuration', {})
            if config.get('openai_api_key'):
                api_key = config['openai_api_key']
                break
        
        if not api_key:
            api_key = os.getenv('OPENAI_API_KEY')
        
        if not api_key:
            raise HTTPException(status_code=503, detail="OpenAI API key not configured")
        
        client = openai.AsyncOpenAI(api_key=api_key)
        response = await client.embeddings.create(
            model=model_name,
            input=text
        )
        return response.data[0].embedding
    
    # Add support for other providers here
    raise HTTPException(status_code=501, detail=f"Provider {model_config['provider']} not yet implemented")

# Pydantic models
class DocumentMetadata(BaseModel):
    title: Optional[str] = None
    source: Optional[str] = None
    document_type: Optional[str] = None
    upload_date: datetime = Field(default_factory=datetime.utcnow)
    file_size: Optional[int] = None
    tags: List[str] = Field(default_factory=list)

class DocumentCreate(BaseModel):
    content: str = Field(..., description="Document content")
    metadata: Optional[DocumentMetadata] = Field(default_factory=DocumentMetadata)

class DocumentResponse(BaseModel):
    id: str
    content: str
    metadata: DocumentMetadata
    chunks_count: int
    indexed_at: datetime

class SearchRequest(BaseModel):
    query: str = Field(..., description="Search query")
    namespace: str = Field(default="default", description="Namespace to search")
    n_results: int = Field(default=5, description="Number of results to return")
    include_content: bool = Field(default=True, description="Include document content")
    similarity_threshold: float = Field(default=0.7, description="Minimum similarity threshold")
    embedding_model: str = Field(default="text-embedding-3-small", description="Embedding model to use")

class SearchResult(BaseModel):
    id: str
    content: str
    similarity: float
    metadata: Dict[str, Any]

class RAGRequest(BaseModel):
    query: str = Field(..., description="User query")
    namespace: str = Field(default="default", description="Namespace to search")
    n_context: int = Field(default=3, description="Number of context documents")
    model: str = Field(default="gpt-4o", description="Model to use for generation")
    max_tokens: int = Field(default=1000, description="Maximum tokens in response")
    embedding_model: str = Field(default="text-embedding-3-small", description="Embedding model for search")

class RAGResponse(BaseModel):
    query: str
    generated_response: str
    context_documents: List[SearchResult]
    model_used: str
    generation_time_ms: float
    timestamp: datetime = Field(default_factory=datetime.utcnow)

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Application lifespan events"""
    global async_engine, async_session_maker
    
    # Startup
    logger.info("Starting RAG service with PGVector...")
    
    try:
        # Initialize database connection
        async_engine = create_async_engine(
            DATABASE_URL,
            echo=False,
            pool_size=20,
            max_overflow=30
        )
        
        # Create async session maker
        async_session_maker = async_sessionmaker(
            async_engine, expire_on_commit=False
        )
        
        # Test database connection
        async with async_engine.begin() as conn:
            await conn.execute(text("SELECT 1"))
        
        logger.info("Connected to PostgreSQL with PGVector")
        
        # Load configuration from database
        await load_rag_configuration()
        
        # Create upload directory
        UPLOAD_DIR.mkdir(exist_ok=True)
        
    except Exception as e:
        logger.error(f"Failed to initialize RAG service: {e}")
        raise
    
    yield
    
    # Shutdown
    logger.info("Shutting down RAG service...")
    if async_engine:
        await async_engine.dispose()

app = FastAPI(
    title="RAG Service",
    description="Retrieval Augmented Generation service with PGVector and configurable models",
    version="1.0.0",
    lifespan=lifespan
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure appropriately for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    
    pgvector_status = "disconnected"
    if async_engine:
        try:
            async with async_engine.begin() as conn:
                await conn.execute(text("SELECT 1"))
            pgvector_status = "connected"
        except Exception:
            pgvector_status = "error"
    
    return {
        "status": "healthy", 
        "service": "rag",
        "version": "1.0.0",
        "features": {
            "pgvector": pgvector_status,
            "document_indexing": True,
            "semantic_search": True,
            "vector_database": True,
            "file_upload": True,
            "configurable_models": True
        },
        "models": {
            "llm_models": list(llm_models_cache.keys()),
            "embedding_models": list(embedding_models_cache.keys()),
            "rag_instances": list(rag_tool_instances_cache.keys())
        }
    }

@app.get("/")
async def root():
    """Root endpoint"""
    return {
        "service": "RAG Service",
        "version": "1.0.0",
        "description": "Retrieval Augmented Generation service with PGVector and configurable models",
        "endpoints": {
            "documents": "/documents",
            "upload": "/documents/upload",
            "search": "/search",
            "generate": "/generate",
            "models": "/models",
            "health": "/health",
            "docs": "/docs"
        }
    }

@app.get("/models")
async def list_models():
    """List available models and configurations"""
    return {
        "llm_models": llm_models_cache,
        "embedding_models": embedding_models_cache,
        "rag_tool_instances": rag_tool_instances_cache
    }

@app.post("/models/reload")
async def reload_models():
    """Reload model configurations from database"""
    try:
        await load_rag_configuration()
        return {
            "message": "Models reloaded successfully",
            "llm_models_count": len(llm_models_cache),
            "embedding_models_count": len(embedding_models_cache),
            "rag_instances_count": len(rag_tool_instances_cache)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to reload models: {str(e)}")

@app.post("/documents/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    namespace: str = Form(default="default"),
    title: Optional[str] = Form(None),
    tags: Optional[str] = Form(None),
    embedding_model: str = Form(default="text-embedding-3-small"),
    session: AsyncSession = Depends(get_database_session)
):
    """Upload and index a document"""
    
    try:
        # Save uploaded file
        file_id = str(uuid.uuid4())
        file_extension = Path(file.filename).suffix.lower() if file.filename else ""
        file_path = UPLOAD_DIR / f"{file_id}{file_extension}"
        
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
        
        # Extract text content based on file type
        if file_extension == ".pdf":
            text_content = extract_pdf_text(file_path)
        elif file_extension in [".docx", ".doc"]:
            text_content = extract_docx_text(file_path)
        elif file_extension == ".txt":
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                text_content = await f.read()
        else:
            # Try to read as text
            try:
                async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                    text_content = await f.read()
            except UnicodeDecodeError:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Unsupported file type: {file_extension}"
                )
        
        # Create metadata
        metadata = DocumentMetadata(
            title=title or file.filename,
            source=file.filename,
            document_type=file_extension.lstrip('.'),
            file_size=len(content),
            tags=tags.split(',') if tags else []
        )
        
        # Index document
        doc_response = await index_document_content(
            session, file_id, text_content, metadata, namespace, embedding_model
        )
        
        return doc_response
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading document: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to upload document: {str(e)}"
        )

@app.post("/documents", response_model=DocumentResponse)
async def index_document(
    document: DocumentCreate,
    namespace: str = "default",
    embedding_model: str = "text-embedding-3-small",
    session: AsyncSession = Depends(get_database_session)
):
    """Index a document from provided content"""
    
    try:
        doc_id = str(uuid.uuid4())
        metadata = document.metadata or DocumentMetadata()
        
        doc_response = await index_document_content(
            session, doc_id, document.content, metadata, namespace, embedding_model
        )
        
        return doc_response
        
    except Exception as e:
        logger.error(f"Error indexing document: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to index document: {str(e)}"
        )

async def index_document_content(
    session: AsyncSession,
    doc_id: str,
    content: str,
    metadata: DocumentMetadata,
    namespace: str,
    embedding_model: str
) -> DocumentResponse:
    """Index document content in PGVector"""
    

    # Use Docling for chunking if possible
    try:
        docling_doc = DoclingDocument(content)
        # You can adjust chunking strategy as needed
        chunks = [chunk.text for chunk in chunk_document(docling_doc, chunk_size=1000, overlap=200)]
        if not chunks:
            # fallback to old chunking if docling fails
            chunks = chunk_text(content, chunk_size=1000, overlap=200)
    except Exception as e:
        logger.warning(f"Docling chunking failed, falling back to legacy: {e}")
        chunks = chunk_text(content, chunk_size=1000, overlap=200)
    
    # Process each chunk
    for i, chunk in enumerate(chunks):
        try:
            # Get embedding for chunk
            embedding = await get_embedding_from_model(chunk, embedding_model)
            
            # Prepare metadata for this chunk
            chunk_metadata = {
                "document_id": doc_id,
                "chunk_index": i,
                "title": metadata.title,
                "source": metadata.source,
                "document_type": metadata.document_type,
                "upload_date": metadata.upload_date.isoformat(),
                "tags": metadata.tags,
                "file_size": metadata.file_size
            }
            
            # Insert into document_embeddings table
            await session.execute(text("""
                INSERT INTO document_embeddings 
                (document_id, content, embedding, metadata, namespace, created_at, updated_at)
                VALUES (:document_id, :content, :embedding, :metadata, :namespace, :created_at, :updated_at)
            """), {
                "document_id": f"{doc_id}_chunk_{i}",
                "content": chunk,
                "embedding": embedding,
                "metadata": json.dumps(chunk_metadata),
                "namespace": namespace,
                "created_at": datetime.utcnow(),
                "updated_at": datetime.utcnow()
            })
            
        except Exception as e:
            logger.error(f"Error processing chunk {i}: {e}")
            continue
    
    await session.commit()
    
    logger.info(f"Indexed document {doc_id} with {len(chunks)} chunks")
    
    return DocumentResponse(
        id=doc_id,
        content=content[:500] + "..." if len(content) > 500 else content,
        metadata=metadata,
        chunks_count=len(chunks),
        indexed_at=datetime.utcnow()
    )

@app.post("/search")
async def semantic_search(
    request: SearchRequest,
    session: AsyncSession = Depends(get_database_session)
):
    """Perform semantic search on indexed documents"""
    
    try:
        # Get embedding for query
        query_embedding = await get_embedding_from_model(request.query, request.embedding_model)
        
        # Search similar documents using PGVector
        result = await session.execute(text("""
            SELECT id, document_id, content, 
                   cosine_similarity(embedding, :query_embedding) as similarity,
                   metadata
            FROM document_embeddings
            WHERE namespace = :namespace
            AND cosine_similarity(embedding, :query_embedding) > :similarity_threshold
            ORDER BY similarity DESC
            LIMIT :limit
        """), {
            "query_embedding": query_embedding,
            "namespace": request.namespace,
            "similarity_threshold": request.similarity_threshold,
            "limit": request.n_results
        })
        
        # Format results
        search_results = []
        for row in result:
            search_results.append(SearchResult(
                id=row[1],  # document_id
                content=row[2] if request.include_content else "",
                similarity=row[3],
                metadata=json.loads(row[4]) if row[4] else {}
            ))
        
        return {
            "query": request.query,
            "namespace": request.namespace,
            "results": search_results,
            "total_results": len(search_results),
            "embedding_model": request.embedding_model,
            "timestamp": datetime.utcnow()
        }
        
    except Exception as e:
        logger.error(f"Error in semantic search: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to perform search: {str(e)}"
        )

@app.post("/generate", response_model=RAGResponse)
async def rag_generate(
    request: RAGRequest,
    session: AsyncSession = Depends(get_database_session)
):
    """Generate response using RAG (Retrieval + Generation)"""
    
    start_time = datetime.utcnow()
    
    try:
        # First, perform semantic search to get context
        search_request = SearchRequest(
            query=request.query,
            namespace=request.namespace,
            n_results=request.n_context,
            include_content=True,
            embedding_model=request.embedding_model
        )
        
        search_response = await semantic_search(search_request, session)
        context_docs = search_response["results"]
        
        # Prepare context for generation
        context_text = "\n\n".join([
            f"Document {i+1}:\n{doc.content}"
            for i, doc in enumerate(context_docs)
        ])
        
        # Get model configuration
        if request.model not in llm_models_cache:
            raise HTTPException(status_code=404, detail=f"LLM model {request.model} not found")
        
        model_config = llm_models_cache[request.model]
        
        # Generate response based on provider
        generated_text = await generate_with_model(
            model_config, context_text, request.query, request.max_tokens
        )
        
        end_time = datetime.utcnow()
        generation_time_ms = (end_time - start_time).total_seconds() * 1000
        
        return RAGResponse(
            query=request.query,
            generated_response=generated_text,
            context_documents=context_docs,
            model_used=request.model,
            generation_time_ms=generation_time_ms
        )
        
    except Exception as e:
        logger.error(f"Error in RAG generation: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate response: {str(e)}"
        )

async def generate_with_model(model_config: Dict, context: str, query: str, max_tokens: int) -> str:
    """Generate text using configured model"""
    
    provider = model_config['provider']
    
    if provider == 'OpenAI':
        # Get API key from tool instance configuration
        api_key = None
        for instance in rag_tool_instances_cache.values():
            config = instance.get('configuration', {})
            if config.get('openai_api_key'):
                api_key = config['openai_api_key']
                break
        
        if not api_key:
            api_key = os.getenv('OPENAI_API_KEY')
        
        if not api_key:
            raise HTTPException(status_code=503, detail="OpenAI API key not configured")
        
        client = openai.AsyncOpenAI(api_key=api_key)
        
        prompt = f"""Based on the following context documents, please answer the user's question. If the answer cannot be found in the context, please say so.

Context:
{context}

Question: {query}

Answer:"""

        response = await client.chat.completions.create(
            model=model_config['name'],
            messages=[{"role": "user", "content": prompt}],
            max_tokens=max_tokens,
            temperature=0.7
        )
        
        return response.choices[0].message.content

    elif provider == 'Google':
        # Get API key from tool instance configuration
        api_key = None
        for instance in rag_tool_instances_cache.values():
            config = instance.get('configuration', {})
            if config.get('google_api_key'):
                api_key = config['google_api_key']
                break
        
        if not api_key:
            api_key = os.getenv('GOOGLE_API_KEY')
        
        if not api_key:
            raise HTTPException(status_code=503, detail="Google API key not configured")
        
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_config['name'])
        
        prompt = f"""Based on the following context documents, please answer the user's question. If the answer cannot be found in the context, please say so.

Context:
{context}

Question: {query}

Answer:"""
        
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                max_output_tokens=max_tokens,
                temperature=0.7
            )
        )
        
        return response.text
    
    else:
        raise HTTPException(status_code=501, detail=f"Provider {provider} not yet implemented")

def extract_pdf_text(file_path: Path) -> str:
    """Extract text from PDF file"""
    
    try:
        reader = pypdf.PdfReader(str(file_path))
        text = ""
        
        for page in reader.pages:
            text += page.extract_text() + "\n"
        
        return text.strip()
        
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        raise Exception(f"Failed to extract PDF text: {e}")

def extract_docx_text(file_path: Path) -> str:
    """Extract text from DOCX file"""
    
    try:
        doc = DocxDocument(str(file_path))
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return text.strip()
        
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        raise Exception(f"Failed to extract DOCX text: {e}")

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks"""
    
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = min(start + chunk_size, len(text))
        
        # Try to break at sentence or word boundary
        if end < len(text):
            # Look for sentence end
            last_period = text.rfind('.', start, end)
            last_question = text.rfind('?', start, end)
            last_exclamation = text.rfind('!', start, end)
            
            sentence_end = max(last_period, last_question, last_exclamation)
            
            if sentence_end > start:
                end = sentence_end + 1
            else:
                # Look for word boundary
                last_space = text.rfind(' ', start, end)
                if last_space > start:
                    end = last_space
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        # Move start position with overlap
        start = max(start + 1, end - overlap)
        
        # Avoid infinite loop
        if start >= len(text):
            break
    
    return chunks

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8004)
