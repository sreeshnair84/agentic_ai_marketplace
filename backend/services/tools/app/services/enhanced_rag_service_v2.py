"""
Enhanced RAG Service with Docling Integration and Langgraph Support
Provides comprehensive RAG functionality with advanced document parsing and vector storage
"""

import asyncio
import asyncpg
import json
import logging
import time
import os
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
import tempfile
import aiofiles
import numpy as np
from datetime import datetime
from dataclasses import dataclass
from enum import Enum

# Document processing with Docling
try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.parsers import DoclingParseOptions
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False
    logging.warning("Docling not available. Install with: pip install docling")

# Embedding and vector storage
import openai
from langchain.embeddings import OpenAIEmbeddings, HuggingFaceEmbeddings
from langchain.text_splitter import (
    RecursiveCharacterTextSplitter, 
    CharacterTextSplitter,
    SemanticSplitter
)
from langchain.vectorstores import FAISS, Chroma, PGVector
from langchain.schema import Document
from langchain.tools import Tool
from langchain.agents import AgentExecutor, create_openai_functions_agent
from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder

# Langgraph integration
try:
    from langgraph import StateGraph, END
    from langgraph.graph import Graph
    from langgraph.checkpoint.sqlite import SqliteSaver
    LANGGRAPH_AVAILABLE = True
except ImportError:
    LANGGRAPH_AVAILABLE = False
    logging.warning("Langgraph not available. Install with: pip install langgraph")

logger = logging.getLogger(__name__)

class DocumentType(Enum):
    """Supported document types"""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    HTML = "html"
    CSV = "csv"
    JSON = "json"
    MARKDOWN = "md"
    XML = "xml"

class ChunkingStrategy(Enum):
    """Available chunking strategies"""
    RECURSIVE = "recursive"
    CHARACTER = "character"
    SEMANTIC = "semantic"
    SENTENCE = "sentence"
    PARAGRAPH = "paragraph"

class EmbeddingProvider(Enum):
    """Supported embedding providers"""
    OPENAI = "openai"
    HUGGINGFACE = "huggingface"
    AZURE_OPENAI = "azure_openai"
    LOCAL = "local"

@dataclass
class RAGConfiguration:
    """RAG Pipeline Configuration"""
    pipeline_id: str
    embedding_model_id: str
    embedding_provider: EmbeddingProvider
    embedding_model_name: str
    vector_database_config: Dict[str, Any]
    chunking_strategy: ChunkingStrategy
    chunk_size: int
    chunk_overlap: int
    use_docling: bool = True
    extract_tables: bool = True
    extract_images: bool = True
    llm_model_id: Optional[str] = None
    quality_threshold: float = 0.7

@dataclass
class ProcessingResult:
    """Document processing result"""
    success: bool
    chunks_created: int = 0
    embeddings_created: int = 0
    tables_extracted: int = 0
    images_extracted: int = 0
    processing_time: float = 0.0
    error_message: Optional[str] = None
    metadata: Dict[str, Any] = None

class EnhancedRAGServiceV2:
    """Enhanced RAG service with Docling integration and Langgraph support"""
    
    def __init__(self, database_url: str = None):
        self.database_url = database_url
        self.connection_pool = None
        self.embedding_cache = {}
        self.vector_store_cache = {}
        self.logger = logger
        
        # Document converter for Docling
        self.document_converter = None
        if DOCLING_AVAILABLE:
            self.document_converter = DocumentConverter()
    
    async def initialize(self, database_url: str = None):
        """Initialize the RAG service with database connection"""
        if database_url:
            self.database_url = database_url
        
        if self.database_url:
            self.connection_pool = await asyncpg.create_pool(self.database_url)
        
        self.logger.info("Enhanced RAG Service V2 initialized")
    
    async def get_rag_configuration(self, pipeline_id: str) -> Optional[RAGConfiguration]:
        """Get RAG configuration from database"""
        if not self.connection_pool:
            raise ValueError("Database connection not initialized")
        
        async with self.connection_pool.acquire() as conn:
            # Get RAG pipeline config
            pipeline_query = """
                SELECT rp.*, ti.llm_model_id, ti.embedding_model_id, ti.configuration
                FROM rag_pipelines rp
                LEFT JOIN tool_instances ti ON ti.id = rp.tool_instance_id
                WHERE rp.id = $1
            """
            pipeline_row = await conn.fetchrow(pipeline_query, pipeline_id)
            
            if not pipeline_row:
                return None
            
            # Get embedding model details
            embedding_query = """
                SELECT * FROM embedding_models WHERE id = $1
            """
            embedding_row = await conn.fetchrow(embedding_query, pipeline_row['embedding_model_id'])
            
            if not embedding_row:
                self.logger.warning(f"Embedding model not found for pipeline {pipeline_id}")
                return None
            
            # Parse configuration
            vectorization_config = pipeline_row['vectorization_config'] or {}
            processing_config = pipeline_row['processing_config'] or {}
            
            return RAGConfiguration(
                pipeline_id=pipeline_id,
                embedding_model_id=str(pipeline_row['embedding_model_id']),
                embedding_provider=EmbeddingProvider(embedding_row['provider']),
                embedding_model_name=embedding_row['name'],
                vector_database_config=pipeline_row['storage_config'] or {},
                chunking_strategy=ChunkingStrategy(vectorization_config.get('text_splitter', 'recursive')),
                chunk_size=vectorization_config.get('chunk_size', 1000),
                chunk_overlap=vectorization_config.get('chunk_overlap', 200),
                use_docling=processing_config.get('use_docling', True),
                extract_tables=processing_config.get('extract_tables', True),
                extract_images=processing_config.get('extract_images', True),
                llm_model_id=str(pipeline_row['llm_model_id']) if pipeline_row['llm_model_id'] else None,
                quality_threshold=processing_config.get('quality_threshold', 0.7)
            )
    
    async def process_document_with_docling(
        self, 
        file_path: str, 
        config: RAGConfiguration
    ) -> ProcessingResult:
        """Process document using Docling with advanced parsing"""
        
        if not DOCLING_AVAILABLE:
            raise RuntimeError("Docling not available. Please install docling package.")
        
        start_time = time.time()
        
        try:
            # Configure Docling pipeline
            pipeline_options = PdfPipelineOptions()
            pipeline_options.do_ocr = True
            pipeline_options.do_table_structure = config.extract_tables
            pipeline_options.table_structure_options.do_cell_matching = True
            
            # Convert document
            result = self.document_converter.convert(file_path, pipeline_options=pipeline_options)
            
            # Extract content
            main_text = result.document.export_to_text()
            
            # Extract tables if configured
            tables = []
            if config.extract_tables:
                for table in result.document.tables:
                    table_data = {
                        'content': table.export_to_text(),
                        'structured_data': table.data if hasattr(table, 'data') else None,
                        'bbox': table.prov[0].bbox if table.prov else None
                    }
                    tables.append(table_data)
            
            # Extract images if configured
            images = []
            if config.extract_images:
                for figure in result.document.pictures:
                    image_data = {
                        'caption': figure.caption if hasattr(figure, 'caption') else '',
                        'bbox': figure.prov[0].bbox if figure.prov else None,
                        'content': f"[Image: {figure.caption if hasattr(figure, 'caption') else 'No caption'}]"
                    }
                    images.append(image_data)
            
            # Create chunks
            chunks = await self._create_chunks_from_content(
                main_text, tables, images, config
            )
            
            processing_time = time.time() - start_time
            
            return ProcessingResult(
                success=True,
                chunks_created=len(chunks),
                tables_extracted=len(tables),
                images_extracted=len(images),
                processing_time=processing_time,
                metadata={
                    'main_text_length': len(main_text),
                    'document_type': 'pdf',
                    'processing_method': 'docling'
                }
            )
            
        except Exception as e:
            self.logger.error(f"Error processing document with Docling: {e}")
            return ProcessingResult(
                success=False,
                error_message=str(e),
                processing_time=time.time() - start_time
            )
    
    async def _create_chunks_from_content(
        self,
        main_text: str,
        tables: List[Dict],
        images: List[Dict],
        config: RAGConfiguration
    ) -> List[Document]:
        """Create chunks from extracted content"""
        
        documents = []
        
        # Create text splitter based on strategy
        if config.chunking_strategy == ChunkingStrategy.RECURSIVE:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=config.chunk_size,
                chunk_overlap=config.chunk_overlap,
                separators=["\n\n", "\n", " ", ""]
            )
        elif config.chunking_strategy == ChunkingStrategy.CHARACTER:
            text_splitter = CharacterTextSplitter(
                chunk_size=config.chunk_size,
                chunk_overlap=config.chunk_overlap,
                separator="\n\n"
            )
        elif config.chunking_strategy == ChunkingStrategy.SEMANTIC:
            # Note: SemanticSplitter requires an embedding model
            embedding_model = await self._get_embedding_model(config)
            text_splitter = SemanticSplitter(
                embeddings=embedding_model,
                buffer_size=1
            )
        else:
            # Default to recursive
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=config.chunk_size,
                chunk_overlap=config.chunk_overlap
            )
        
        # Split main text
        main_chunks = text_splitter.split_text(main_text)
        for i, chunk in enumerate(main_chunks):
            documents.append(Document(
                page_content=chunk,
                metadata={
                    'chunk_index': i,
                    'content_type': 'text',
                    'pipeline_id': config.pipeline_id,
                    'created_at': datetime.utcnow().isoformat()
                }
            ))
        
        # Add table chunks
        for i, table in enumerate(tables):
            table_chunks = text_splitter.split_text(table['content'])
            for j, chunk in enumerate(table_chunks):
                documents.append(Document(
                    page_content=chunk,
                    metadata={
                        'chunk_index': f"table_{i}_{j}",
                        'content_type': 'table',
                        'table_index': i,
                        'pipeline_id': config.pipeline_id,
                        'bbox': table.get('bbox'),
                        'created_at': datetime.utcnow().isoformat()
                    }
                ))
        
        # Add image descriptions as chunks
        for i, image in enumerate(images):
            if image['content']:
                documents.append(Document(
                    page_content=image['content'],
                    metadata={
                        'chunk_index': f"image_{i}",
                        'content_type': 'image',
                        'image_index': i,
                        'pipeline_id': config.pipeline_id,
                        'bbox': image.get('bbox'),
                        'created_at': datetime.utcnow().isoformat()
                    }
                ))
        
        return documents
    
    async def _get_embedding_model(self, config: RAGConfiguration):
        """Get embedding model based on configuration"""
        
        if config.embedding_model_id in self.embedding_cache:
            return self.embedding_cache[config.embedding_model_id]
        
        # Get embedding model details from database
        async with self.connection_pool.acquire() as conn:
            query = "SELECT * FROM embedding_models WHERE id = $1"
            row = await conn.fetchrow(query, config.embedding_model_id)
            
            if not row:
                raise ValueError(f"Embedding model {config.embedding_model_id} not found")
            
            # Create embedding model based on provider
            if config.embedding_provider == EmbeddingProvider.OPENAI:
                embeddings = OpenAIEmbeddings(
                    model=row['name'],
                    openai_api_key=os.getenv(row['api_key_env_var']),
                    openai_api_base=row['endpoint_url']
                )
            elif config.embedding_provider == EmbeddingProvider.HUGGINGFACE:
                embeddings = HuggingFaceEmbeddings(
                    model_name=row['name']
                )
            else:
                raise ValueError(f"Unsupported embedding provider: {config.embedding_provider}")
            
            self.embedding_cache[config.embedding_model_id] = embeddings
            return embeddings
    
    async def create_rag_tools_for_langgraph(self, pipeline_id: str) -> List[Tool]:
        """Create Langchain tools for RAG operations that can be used with Langgraph"""
        
        config = await self.get_rag_configuration(pipeline_id)
        if not config:
            raise ValueError(f"RAG pipeline {pipeline_id} not found")
        
        tools = []
        
        # Search tool
        async def rag_search(query: str, k: int = 5) -> str:
            """Search the RAG knowledge base for relevant information"""
            try:
                results = await self.search_pipeline(
                    pipeline_id=pipeline_id,
                    query=query,
                    k=k
                )
                
                if not results:
                    return "No relevant documents found."
                
                formatted_results = []
                for i, result in enumerate(results[:k]):
                    formatted_results.append(
                        f"Result {i+1}:\n"
                        f"Content: {result['content']}\n"
                        f"Source: {result.get('metadata', {}).get('filename', 'Unknown')}\n"
                        f"Relevance: {result.get('score', 0):.3f}\n"
                    )
                
                return "\n".join(formatted_results)
                
            except Exception as e:
                return f"Error searching knowledge base: {str(e)}"
        
        tools.append(Tool(
            name="rag_search",
            description=f"Search the RAG knowledge base (Pipeline: {pipeline_id}) for relevant documents and information",
            func=rag_search
        ))
        
        # Document upload tool
        async def upload_document(file_path: str, metadata: str = "{}") -> str:
            """Upload and process a document into the RAG knowledge base"""
            try:
                metadata_dict = json.loads(metadata) if metadata else {}
                
                result = await self.ingest_document(
                    pipeline_id=pipeline_id,
                    file_path=file_path,
                    filename=Path(file_path).name,
                    metadata=metadata_dict
                )
                
                if result.success:
                    return f"Successfully processed document. Created {result.chunks_created} chunks, extracted {result.tables_extracted} tables and {result.images_extracted} images."
                else:
                    return f"Failed to process document: {result.error_message}"
                    
            except Exception as e:
                return f"Error uploading document: {str(e)}"
        
        tools.append(Tool(
            name="upload_document",
            description=f"Upload and process a document into the RAG knowledge base (Pipeline: {pipeline_id})",
            func=upload_document
        ))
        
        return tools
    
    async def create_langgraph_rag_agent(self, pipeline_id: str, llm_model=None):
        """Create a Langgraph agent with RAG capabilities"""
        
        if not LANGGRAPH_AVAILABLE:
            raise RuntimeError("Langgraph not available. Please install langgraph package.")
        
        config = await self.get_rag_configuration(pipeline_id)
        if not config:
            raise ValueError(f"RAG pipeline {pipeline_id} not found")
        
        # Get RAG tools
        tools = await self.create_rag_tools_for_langgraph(pipeline_id)
        
        # Define the agent state
        class AgentState:
            def __init__(self):
                self.messages: List[Dict] = []
                self.intermediate_steps: List[Dict] = []
                self.final_response: str = ""
        
        # Create the graph
        graph = StateGraph(AgentState)
        
        # RAG Search Node
        async def rag_search_node(state: AgentState):
            """Node for performing RAG search"""
            last_message = state.messages[-1] if state.messages else {"content": ""}
            query = last_message.get("content", "")
            
            if not query:
                state.final_response = "No query provided"
                return state
            
            # Perform RAG search
            try:
                results = await self.search_pipeline(pipeline_id, query)
                
                context = "\n\n".join([
                    f"Document: {r.get('metadata', {}).get('filename', 'Unknown')}\n"
                    f"Content: {r['content'][:500]}..."
                    for r in results[:3]
                ])
                
                # Generate response using context
                if llm_model:
                    prompt = f"""
Based on the following context from the knowledge base, please provide a helpful response to the query: "{query}"

Context:
{context}

Query: {query}

Response:"""
                    response = await llm_model.agenerate([prompt])
                    state.final_response = response.generations[0][0].text
                else:
                    state.final_response = f"Found {len(results)} relevant documents:\n\n{context}"
                
            except Exception as e:
                state.final_response = f"Error performing RAG search: {str(e)}"
            
            return state
        
        # Add nodes to graph
        graph.add_node("rag_search", rag_search_node)
        graph.add_edge("rag_search", END)
        graph.set_entry_point("rag_search")
        
        # Compile the graph
        app = graph.compile()
        
        return app
    
    async def search_pipeline(
        self, 
        pipeline_id: str, 
        query: str, 
        k: int = 5, 
        filters: Dict[str, Any] = None
    ) -> List[Dict[str, Any]]:
        """Search documents in a RAG pipeline"""
        
        config = await self.get_rag_configuration(pipeline_id)
        if not config:
            raise ValueError(f"RAG pipeline {pipeline_id} not found")
        
        try:
            # Get vector store
            vector_store = await self._get_vector_store(config)
            
            # Perform search
            results = vector_store.similarity_search_with_score(
                query=query,
                k=k,
                filter=filters
            )
            
            # Format results
            formatted_results = []
            for doc, score in results:
                formatted_results.append({
                    'content': doc.page_content,
                    'metadata': doc.metadata,
                    'score': float(score)
                })
            
            return formatted_results
            
        except Exception as e:
            self.logger.error(f"Error searching pipeline {pipeline_id}: {e}")
            return []
    
    async def ingest_document(
        self,
        pipeline_id: str,
        file_path: str,
        filename: str,
        metadata: Dict[str, Any] = None
    ) -> ProcessingResult:
        """Ingest a document into the RAG pipeline"""
        
        config = await self.get_rag_configuration(pipeline_id)
        if not config:
            raise ValueError(f"RAG pipeline {pipeline_id} not found")
        
        # Use Docling for advanced processing if available and configured
        if config.use_docling and DOCLING_AVAILABLE:
            return await self.process_document_with_docling(file_path, config)
        else:
            # Fallback to basic processing
            return await self._process_document_basic(file_path, filename, config, metadata)
    
    async def _process_document_basic(
        self,
        file_path: str,
        filename: str,
        config: RAGConfiguration,
        metadata: Dict[str, Any] = None
    ) -> ProcessingResult:
        """Basic document processing without Docling"""
        
        start_time = time.time()
        
        try:
            # Extract text content
            text_content = await self._extract_text_basic(file_path, filename)
            
            if not text_content.strip():
                return ProcessingResult(
                    success=False,
                    error_message="No text content extracted from document"
                )
            
            # Create chunks
            documents = await self._create_chunks_from_content(
                text_content, [], [], config
            )
            
            # Add to vector store
            vector_store = await self._get_vector_store(config)
            await vector_store.aadd_documents(documents)
            
            processing_time = time.time() - start_time
            
            return ProcessingResult(
                success=True,
                chunks_created=len(documents),
                processing_time=processing_time,
                metadata={'processing_method': 'basic'}
            )
            
        except Exception as e:
            self.logger.error(f"Error processing document {filename}: {e}")
            return ProcessingResult(
                success=False,
                error_message=str(e),
                processing_time=time.time() - start_time
            )
    
    async def _extract_text_basic(self, file_path: str, filename: str) -> str:
        """Basic text extraction from common file formats"""
        
        file_ext = Path(filename).suffix.lower()
        
        if file_ext == '.pdf':
            return await self._extract_from_pdf(file_path)
        elif file_ext == '.docx':
            return await self._extract_from_docx(file_path)
        elif file_ext in ['.txt', '.md']:
            return await self._extract_from_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        try:
            import pypdf
            with open(file_path, 'rb') as file:
                reader = pypdf.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except ImportError:
            raise RuntimeError("pypdf not available. Install with: pip install pypdf")
    
    async def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except ImportError:
            raise RuntimeError("python-docx not available. Install with: pip install python-docx")
    
    async def _extract_from_text(self, file_path: str) -> str:
        """Extract text from plain text files"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            return await file.read()
    
    async def _get_vector_store(self, config: RAGConfiguration):
        """Get or create vector store for the pipeline"""
        
        if config.pipeline_id in self.vector_store_cache:
            return self.vector_store_cache[config.pipeline_id]
        
        embeddings = await self._get_embedding_model(config)
        
        # Create vector store based on configuration
        vector_db_config = config.vector_database_config
        
        if vector_db_config.get('type') == 'pgvector':
            vector_store = PGVector(
                connection_string=vector_db_config.get('connection_string'),
                embedding_function=embeddings,
                collection_name=f"pipeline_{config.pipeline_id}"
            )
        elif vector_db_config.get('type') == 'chroma':
            vector_store = Chroma(
                collection_name=f"pipeline_{config.pipeline_id}",
                embedding_function=embeddings,
                persist_directory=vector_db_config.get('persist_directory', './chroma_db')
            )
        else:
            # Default to FAISS
            vector_store = FAISS(
                embedding_function=embeddings,
                index=None,
                docstore={},
                index_to_docstore_id={}
            )
        
        self.vector_store_cache[config.pipeline_id] = vector_store
        return vector_store
    
    async def create_mcp_endpoint(self, pipeline_id: str) -> Dict[str, Any]:
        """Create MCP endpoint configuration for the RAG pipeline"""
        
        config = await self.get_rag_configuration(pipeline_id)
        if not config:
            raise ValueError(f"RAG pipeline {pipeline_id} not found")
        
        # Create tools for MCP
        tools = await self.create_rag_tools_for_langgraph(pipeline_id)
        
        mcp_config = {
            "name": f"rag_pipeline_{pipeline_id}",
            "description": f"RAG Pipeline {pipeline_id} - Document search and knowledge base access",
            "version": "1.0.0",
            "tools": [
                {
                    "name": tool.name,
                    "description": tool.description,
                    "input_schema": {
                        "type": "object",
                        "properties": {
                            "query": {"type": "string", "description": "Search query"},
                            "k": {"type": "integer", "description": "Number of results", "default": 5}
                        }
                    }
                }
                for tool in tools
            ],
            "endpoints": {
                "search": f"/rag/{pipeline_id}/search",
                "upload": f"/rag/{pipeline_id}/upload",
                "health": f"/rag/{pipeline_id}/health"
            }
        }
        
        return mcp_config
    
    async def health_check(self, pipeline_id: str) -> Dict[str, Any]:
        """Perform health check on RAG pipeline"""
        
        try:
            config = await self.get_rag_configuration(pipeline_id)
            if not config:
                return {
                    "status": "error",
                    "message": "Pipeline not found"
                }
            
            # Check embedding model
            try:
                embedding_model = await self._get_embedding_model(config)
                embedding_test = await embedding_model.aembed_query("test")
                embedding_status = "healthy" if embedding_test else "error"
            except Exception as e:
                embedding_status = f"error: {str(e)}"
            
            # Check vector store
            try:
                vector_store = await self._get_vector_store(config)
                vector_status = "healthy"
            except Exception as e:
                vector_status = f"error: {str(e)}"
            
            return {
                "status": "healthy" if embedding_status == "healthy" and vector_status == "healthy" else "degraded",
                "pipeline_id": pipeline_id,
                "embedding_model_status": embedding_status,
                "vector_store_status": vector_status,
                "docling_available": DOCLING_AVAILABLE,
                "langgraph_available": LANGGRAPH_AVAILABLE,
                "timestamp": datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            return {
                "status": "error",
                "message": str(e),
                "timestamp": datetime.utcnow().isoformat()
            }