"""
PostgreSQL Enhanced RAG Tool Implementation
Advanced RAG system with PostgreSQL vector database integration using pgvector
"""

import asyncio
import asyncpg
import json
import logging
import time
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
from pathlib import Path
import uuid
import hashlib
import aiofiles
import tempfile

# Document processing
try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False

# Vector and embedding support
import numpy as np
from pgvector.asyncpg import register_vector
import openai
from sentence_transformers import SentenceTransformer

# Text processing
from langchain.text_splitter import (
    RecursiveCharacterTextSplitter,
    CharacterTextSplitter,
    SemanticSplitter
)
from langchain.embeddings import OpenAIEmbeddings, HuggingFaceEmbeddings

logger = logging.getLogger(__name__)

class PostgresEnhancedRAGTool:
    """
    Enhanced RAG implementation using PostgreSQL with pgvector extension
    Features advanced document processing, semantic search, and comprehensive metadata management
    """
    
    def __init__(self, config: Dict[str, Any]):
        """Initialize PostgreSQL RAG tool with configuration"""
        self.config = config
        
        # Database configuration
        self.database_url = config.get("database_url", "postgresql://user:pass@localhost/dbname")
        self.table_name = config.get("table_name", "rag_documents")
        
        # RAG configuration
        self.embedding_model = config.get("embedding_model", "text-embedding-3-small")
        self.embedding_provider = config.get("embedding_provider", "openai")
        self.chunk_size = config.get("chunk_size", 1000)
        self.chunk_overlap = config.get("chunk_overlap", 200)
        self.chunking_strategy = config.get("chunking_strategy", "recursive")
        
        # Advanced processing options
        self.use_docling = config.get("use_docling", True)
        self.extract_tables = config.get("extract_tables", True)
        self.extract_images = config.get("extract_images", True)
        self.quality_threshold = config.get("quality_threshold", 0.7)
        
        # Initialize components
        self.connection_pool = None
        self.embedding_model_instance = None
        self.document_converter = None
        self.text_splitter = None
        
        if DOCLING_AVAILABLE and self.use_docling:
            self.document_converter = DocumentConverter()
        
        logger.info(f"Initialized PostgreSQL RAG tool with embedding model: {self.embedding_model}")
    
    async def initialize(self):
        """Initialize database connection and embedding model"""
        try:
            # Create connection pool
            self.connection_pool = await asyncpg.create_pool(self.database_url)
            
            # Register pgvector types
            async with self.connection_pool.acquire() as conn:
                await register_vector(conn)
            
            # Initialize embedding model
            await self._initialize_embedding_model()
            
            # Initialize text splitter
            self._initialize_text_splitter()
            
            # Create database tables if they don't exist
            await self._create_tables()
            
            logger.info("PostgreSQL RAG tool initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize PostgreSQL RAG tool: {e}")
            raise
    
    async def _initialize_embedding_model(self):
        """Initialize the embedding model based on provider"""
        if self.embedding_provider == "openai":
            self.embedding_model_instance = OpenAIEmbeddings(
                model=self.embedding_model,
                openai_api_key=self.config.get("openai_api_key")
            )
        elif self.embedding_provider == "huggingface":
            self.embedding_model_instance = HuggingFaceEmbeddings(
                model_name=self.embedding_model
            )
        else:
            raise ValueError(f"Unsupported embedding provider: {self.embedding_provider}")
    
    def _initialize_text_splitter(self):
        """Initialize text splitter based on strategy"""
        if self.chunking_strategy == "recursive":
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separators=["\n\n", "\n", " ", ""]
            )
        elif self.chunking_strategy == "character":
            self.text_splitter = CharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separator="\n\n"
            )
        elif self.chunking_strategy == "semantic" and self.embedding_model_instance:
            self.text_splitter = SemanticSplitter(
                embeddings=self.embedding_model_instance,
                buffer_size=1
            )
        else:
            # Default to recursive
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap
            )
    
    async def _create_tables(self):
        """Create necessary database tables with pgvector support"""
        async with self.connection_pool.acquire() as conn:
            # Enable pgvector extension
            await conn.execute("CREATE EXTENSION IF NOT EXISTS vector")
            
            # Create main documents table
            await conn.execute(f"""
                CREATE TABLE IF NOT EXISTS {self.table_name} (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    document_id VARCHAR(255) NOT NULL,
                    chunk_index INTEGER NOT NULL,
                    content TEXT NOT NULL,
                    content_type VARCHAR(50) DEFAULT 'text',
                    embedding VECTOR(1536),
                    metadata JSONB DEFAULT '{{}}',
                    document_hash VARCHAR(64),
                    file_path TEXT,
                    filename VARCHAR(255),
                    file_size INTEGER,
                    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
                    updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
                    UNIQUE(document_id, chunk_index)
                )
            """)
            
            # Create indexes for performance
            await conn.execute(f"""
                CREATE INDEX IF NOT EXISTS idx_{self.table_name}_embedding 
                ON {self.table_name} USING ivfflat (embedding vector_cosine_ops) 
                WITH (lists = 100)
            """)
            
            await conn.execute(f"""
                CREATE INDEX IF NOT EXISTS idx_{self.table_name}_document_id 
                ON {self.table_name} (document_id)
            """)
            
            await conn.execute(f"""
                CREATE INDEX IF NOT EXISTS idx_{self.table_name}_content_type 
                ON {self.table_name} (content_type)
            """)
            
            await conn.execute(f"""
                CREATE INDEX IF NOT EXISTS idx_{self.table_name}_metadata 
                ON {self.table_name} USING gin (metadata)
            """)
            
            # Create table for document metadata
            await conn.execute(f"""
                CREATE TABLE IF NOT EXISTS {self.table_name}_metadata (
                    document_id VARCHAR(255) PRIMARY KEY,
                    original_filename VARCHAR(255),
                    file_path TEXT,
                    file_size INTEGER,
                    content_type VARCHAR(50),
                    total_chunks INTEGER DEFAULT 0,
                    processing_config JSONB DEFAULT '{{}}',
                    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
                    updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
                )
            """)
    
    async def ingest_document(self, file_path: str, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Ingest a document with enhanced processing capabilities
        
        Args:
            file_path: Path to the document file
            metadata: Additional metadata for the document
            
        Returns:
            Dictionary with ingestion results
        """
        start_time = time.time()
        
        try:
            # Generate document ID
            document_id = str(uuid.uuid4())
            filename = Path(file_path).name
            
            # Calculate file hash for deduplication
            file_hash = await self._calculate_file_hash(file_path)
            
            # Check if document already exists
            existing_doc = await self._check_existing_document(file_hash)
            if existing_doc:
                return {
                    "status": "duplicate",
                    "document_id": existing_doc["document_id"],
                    "message": "Document already exists",
                    "existing_chunks": existing_doc["total_chunks"]
                }
            
            # Process document based on capabilities
            if self.use_docling and DOCLING_AVAILABLE:
                processing_result = await self._process_with_docling(file_path, document_id)
            else:
                processing_result = await self._process_basic(file_path, document_id)
            
            # Store document metadata
            await self._store_document_metadata(
                document_id=document_id,
                file_path=file_path,
                filename=filename,
                file_hash=file_hash,
                total_chunks=processing_result["chunks_created"],
                metadata=metadata or {}
            )
            
            processing_time = time.time() - start_time
            
            return {
                "status": "success",
                "document_id": document_id,
                "filename": filename,
                "chunks_created": processing_result["chunks_created"],
                "tables_extracted": processing_result.get("tables_extracted", 0),
                "images_extracted": processing_result.get("images_extracted", 0),
                "processing_time": processing_time,
                "embedding_model": self.embedding_model,
                "processing_method": processing_result.get("method", "basic")
            }
            
        except Exception as e:
            logger.error(f"Error ingesting document {file_path}: {e}")
            return {
                "status": "error",
                "error": str(e),
                "processing_time": time.time() - start_time
            }
    
    async def _process_with_docling(self, file_path: str, document_id: str) -> Dict[str, Any]:
        """Process document using Docling for advanced extraction"""
        try:
            # Configure Docling pipeline
            pipeline_options = PdfPipelineOptions()
            pipeline_options.do_ocr = True
            pipeline_options.do_table_structure = self.extract_tables
            pipeline_options.table_structure_options.do_cell_matching = True
            
            # Convert document
            result = self.document_converter.convert(file_path, pipeline_options=pipeline_options)
            
            # Extract main content
            main_text = result.document.export_to_text()
            chunks_created = 0
            tables_extracted = 0
            images_extracted = 0
            
            # Process main text
            if main_text.strip():
                text_chunks = self.text_splitter.split_text(main_text)
                for i, chunk in enumerate(text_chunks):
                    await self._store_chunk(
                        document_id=document_id,
                        chunk_index=f"text_{i}",
                        content=chunk,
                        content_type="text",
                        metadata={"source": "main_text", "chunk_type": "text"}
                    )
                    chunks_created += 1
            
            # Process tables if enabled
            if self.extract_tables:
                for table_idx, table in enumerate(result.document.tables):
                    table_content = table.export_to_text()
                    if table_content.strip():
                        await self._store_chunk(
                            document_id=document_id,
                            chunk_index=f"table_{table_idx}",
                            content=table_content,
                            content_type="table",
                            metadata={
                                "source": "table_extraction",
                                "table_index": table_idx,
                                "bbox": table.prov[0].bbox if table.prov else None
                            }
                        )
                        tables_extracted += 1
                        chunks_created += 1
            
            # Process images if enabled
            if self.extract_images:
                for img_idx, figure in enumerate(result.document.pictures):
                    caption = figure.caption if hasattr(figure, 'caption') else f"Image {img_idx + 1}"
                    image_description = f"[Image: {caption}]"
                    
                    await self._store_chunk(
                        document_id=document_id,
                        chunk_index=f"image_{img_idx}",
                        content=image_description,
                        content_type="image",
                        metadata={
                            "source": "image_extraction",
                            "image_index": img_idx,
                            "caption": caption,
                            "bbox": figure.prov[0].bbox if figure.prov else None
                        }
                    )
                    images_extracted += 1
                    chunks_created += 1
            
            return {
                "chunks_created": chunks_created,
                "tables_extracted": tables_extracted,
                "images_extracted": images_extracted,
                "method": "docling"
            }
            
        except Exception as e:
            logger.error(f"Error processing document with Docling: {e}")
            raise
    
    async def _process_basic(self, file_path: str, document_id: str) -> Dict[str, Any]:
        """Basic document processing without Docling"""
        try:
            # Extract text content
            text_content = await self._extract_text_basic(file_path)
            
            if not text_content.strip():
                raise ValueError("No text content extracted from document")
            
            # Create chunks
            chunks = self.text_splitter.split_text(text_content)
            
            # Store chunks
            for i, chunk in enumerate(chunks):
                await self._store_chunk(
                    document_id=document_id,
                    chunk_index=i,
                    content=chunk,
                    content_type="text",
                    metadata={"source": "basic_extraction", "chunk_type": "text"}
                )
            
            return {
                "chunks_created": len(chunks),
                "method": "basic"
            }
            
        except Exception as e:
            logger.error(f"Error in basic document processing: {e}")
            raise
    
    async def _store_chunk(self, document_id: str, chunk_index: Union[int, str], content: str, 
                          content_type: str = "text", metadata: Dict[str, Any] = None):
        """Store a chunk with its embedding in PostgreSQL"""
        
        # Generate embedding
        if self.embedding_provider == "openai":
            embedding_result = await self.embedding_model_instance.aembed_query(content)
            embedding = embedding_result
        else:
            embedding = await self.embedding_model_instance.aembed_query(content)
        
        # Store in database
        async with self.connection_pool.acquire() as conn:
            await conn.execute(f"""
                INSERT INTO {self.table_name} 
                (document_id, chunk_index, content, content_type, embedding, metadata)
                VALUES ($1, $2, $3, $4, $5, $6)
                ON CONFLICT (document_id, chunk_index) 
                DO UPDATE SET 
                    content = EXCLUDED.content,
                    embedding = EXCLUDED.embedding,
                    metadata = EXCLUDED.metadata,
                    updated_at = NOW()
            """, document_id, str(chunk_index), content, content_type, embedding, 
            json.dumps(metadata or {}))
    
    async def search(self, query: str, top_k: int = 5, 
                    content_types: List[str] = None, 
                    filters: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Semantic search with advanced filtering
        
        Args:
            query: Search query
            top_k: Number of results to return
            content_types: Filter by content types (text, table, image)
            filters: Additional metadata filters
            
        Returns:
            Dictionary with search results
        """
        try:
            start_time = time.time()
            
            # Generate query embedding
            if self.embedding_provider == "openai":
                query_embedding = await self.embedding_model_instance.aembed_query(query)
            else:
                query_embedding = await self.embedding_model_instance.aembed_query(query)
            
            # Build query conditions
            conditions = []
            params = [query_embedding, top_k]
            param_idx = 3
            
            if content_types:
                conditions.append(f"content_type = ANY(${param_idx})")
                params.append(content_types)
                param_idx += 1
            
            if filters:
                for key, value in filters.items():
                    conditions.append(f"metadata->>${param_idx} = ${param_idx + 1}")
                    params.extend([key, json.dumps(value) if not isinstance(value, str) else value])
                    param_idx += 2
            
            where_clause = f"WHERE {' AND '.join(conditions)}" if conditions else ""
            
            # Execute search query
            async with self.connection_pool.acquire() as conn:
                query_sql = f"""
                    SELECT 
                        id, document_id, chunk_index, content, content_type, metadata,
                        1 - (embedding <=> $1) as similarity_score,
                        dm.original_filename, dm.created_at as document_created_at
                    FROM {self.table_name} d
                    LEFT JOIN {self.table_name}_metadata dm ON d.document_id = dm.document_id
                    {where_clause}
                    ORDER BY embedding <=> $1
                    LIMIT $2
                """
                
                rows = await conn.fetch(query_sql, *params)
                
                # Format results
                results = []
                for row in rows:
                    result = {
                        "chunk_id": str(row["id"]),
                        "document_id": row["document_id"],
                        "chunk_index": row["chunk_index"],
                        "content": row["content"],
                        "content_type": row["content_type"],
                        "similarity_score": float(row["similarity_score"]),
                        "metadata": json.loads(row["metadata"]) if row["metadata"] else {},
                        "source_info": {
                            "filename": row["original_filename"],
                            "document_created_at": row["document_created_at"].isoformat() if row["document_created_at"] else None
                        }
                    }
                    results.append(result)
            
            search_time = time.time() - start_time
            
            return {
                "status": "success",
                "query": query,
                "total_results": len(results),
                "results": results,
                "search_time": search_time,
                "search_metadata": {
                    "embedding_model": self.embedding_model,
                    "content_types": content_types,
                    "filters": filters
                }
            }
            
        except Exception as e:
            logger.error(f"Error in search: {e}")
            return {
                "status": "error",
                "error": str(e)
            }
    
    async def get_document_statistics(self, document_id: str = None) -> Dict[str, Any]:
        """Get statistics for documents in the RAG system"""
        try:
            async with self.connection_pool.acquire() as conn:
                if document_id:
                    # Statistics for specific document
                    stats = await conn.fetchrow(f"""
                        SELECT 
                            COUNT(*) as total_chunks,
                            COUNT(DISTINCT content_type) as content_types,
                            AVG(length(content)) as avg_chunk_length,
                            MAX(created_at) as last_updated
                        FROM {self.table_name}
                        WHERE document_id = $1
                    """, document_id)
                    
                    content_breakdown = await conn.fetch(f"""
                        SELECT content_type, COUNT(*) as count
                        FROM {self.table_name}
                        WHERE document_id = $1
                        GROUP BY content_type
                    """, document_id)
                else:
                    # Overall statistics
                    stats = await conn.fetchrow(f"""
                        SELECT 
                            COUNT(*) as total_chunks,
                            COUNT(DISTINCT document_id) as total_documents,
                            COUNT(DISTINCT content_type) as content_types,
                            AVG(length(content)) as avg_chunk_length
                        FROM {self.table_name}
                    """)
                    
                    content_breakdown = await conn.fetch(f"""
                        SELECT content_type, COUNT(*) as count
                        FROM {self.table_name}
                        GROUP BY content_type
                    """)
                
                return {
                    "status": "success",
                    "statistics": dict(stats) if stats else {},
                    "content_breakdown": {row["content_type"]: row["count"] for row in content_breakdown},
                    "system_info": {
                        "embedding_model": self.embedding_model,
                        "chunking_strategy": self.chunking_strategy,
                        "docling_enabled": self.use_docling and DOCLING_AVAILABLE
                    }
                }
                
        except Exception as e:
            logger.error(f"Error getting statistics: {e}")
            return {
                "status": "error",
                "error": str(e)
            }
    
    async def delete_document(self, document_id: str) -> Dict[str, Any]:
        """Delete a document and all its chunks"""
        try:
            async with self.connection_pool.acquire() as conn:
                # Delete chunks
                chunks_deleted = await conn.execute(
                    f"DELETE FROM {self.table_name} WHERE document_id = $1",
                    document_id
                )
                
                # Delete metadata
                metadata_deleted = await conn.execute(
                    f"DELETE FROM {self.table_name}_metadata WHERE document_id = $1",
                    document_id
                )
                
                return {
                    "status": "success",
                    "document_id": document_id,
                    "chunks_deleted": int(chunks_deleted.split()[-1]),
                    "metadata_deleted": bool(metadata_deleted)
                }
                
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {e}")
            return {
                "status": "error",
                "error": str(e)
            }
    
    # Utility methods
    
    async def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file"""
        hash_sha256 = hashlib.sha256()
        async with aiofiles.open(file_path, 'rb') as f:
            async for chunk in f:
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    async def _check_existing_document(self, file_hash: str) -> Optional[Dict[str, Any]]:
        """Check if document with same hash already exists"""
        async with self.connection_pool.acquire() as conn:
            row = await conn.fetchrow(f"""
                SELECT document_id, total_chunks 
                FROM {self.table_name}_metadata 
                WHERE document_hash = $1
            """, file_hash)
            
            return dict(row) if row else None
    
    async def _store_document_metadata(self, document_id: str, file_path: str, filename: str,
                                     file_hash: str, total_chunks: int, metadata: Dict[str, Any]):
        """Store document metadata"""
        async with self.connection_pool.acquire() as conn:
            await conn.execute(f"""
                INSERT INTO {self.table_name}_metadata 
                (document_id, original_filename, file_path, file_size, total_chunks, 
                 processing_config, document_hash)
                VALUES ($1, $2, $3, $4, $5, $6, $7)
            """, document_id, filename, file_path, 
            Path(file_path).stat().st_size if Path(file_path).exists() else 0,
            total_chunks, json.dumps(metadata), file_hash)
    
    async def _extract_text_basic(self, file_path: str) -> str:
        """Basic text extraction from common file formats"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return await self._extract_from_pdf(file_path)
        elif file_ext == '.docx':
            return await self._extract_from_docx(file_path)
        elif file_ext in ['.txt', '.md']:
            return await self._extract_from_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        try:
            import pypdf
            with open(file_path, 'rb') as file:
                reader = pypdf.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except ImportError:
            raise RuntimeError("pypdf not available. Install with: pip install pypdf")
    
    async def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except ImportError:
            raise RuntimeError("python-docx not available. Install with: pip install python-docx")
    
    async def _extract_from_text(self, file_path: str) -> str:
        """Extract text from plain text files"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            return await file.read()

# Tool configuration schema
POSTGRES_RAG_CONFIG_SCHEMA = {
    "type": "object",
    "properties": {
        "database_url": {
            "type": "string",
            "description": "PostgreSQL connection URL with pgvector support",
            "default": "postgresql://user:pass@localhost/dbname"
        },
        "table_name": {
            "type": "string",
            "description": "Base table name for RAG data",
            "default": "rag_documents"
        },
        "embedding_model": {
            "type": "string",
            "enum": [
                "text-embedding-3-small",
                "text-embedding-3-large",
                "text-embedding-ada-002",
                "all-MiniLM-L6-v2",
                "all-mpnet-base-v2"
            ],
            "default": "text-embedding-3-small",
            "description": "Embedding model for vectorization"
        },
        "embedding_provider": {
            "type": "string",
            "enum": ["openai", "huggingface"],
            "default": "openai",
            "description": "Embedding model provider"
        },
        "chunk_size": {
            "type": "integer",
            "minimum": 100,
            "maximum": 8000,
            "default": 1000,
            "description": "Size of text chunks in characters"
        },
        "chunk_overlap": {
            "type": "integer",
            "minimum": 0,
            "maximum": 1000,
            "default": 200,
            "description": "Overlap between chunks in characters"
        },
        "chunking_strategy": {
            "type": "string",
            "enum": ["recursive", "character", "semantic"],
            "default": "recursive",
            "description": "Text chunking strategy"
        },
        "use_docling": {
            "type": "boolean",
            "default": True,
            "description": "Use Docling for advanced document processing"
        },
        "extract_tables": {
            "type": "boolean",
            "default": True,
            "description": "Extract and process table content"
        },
        "extract_images": {
            "type": "boolean",
            "default": True,
            "description": "Extract and process image descriptions"
        },
        "quality_threshold": {
            "type": "number",
            "minimum": 0.0,
            "maximum": 1.0,
            "default": 0.7,
            "description": "Quality threshold for content filtering"
        },
        "openai_api_key": {
            "type": "string",
            "description": "OpenAI API key for embeddings (if using OpenAI provider)"
        }
    },
    "required": ["database_url"]
}

# Tool input/output schemas
POSTGRES_RAG_INPUT_SCHEMA = {
    "type": "object",
    "properties": {
        "operation": {
            "type": "string",
            "enum": [
                "ingest_document", 
                "search", 
                "get_statistics", 
                "delete_document",
                "health_check"
            ],
            "description": "Operation to perform"
        },
        "file_path": {
            "type": "string",
            "description": "Path to document file (for ingest_document)"
        },
        "query": {
            "type": "string", 
            "description": "Search query (for search)"
        },
        "top_k": {
            "type": "integer",
            "minimum": 1,
            "maximum": 100,
            "default": 5,
            "description": "Number of search results to return"
        },
        "content_types": {
            "type": "array",
            "items": {"type": "string", "enum": ["text", "table", "image"]},
            "description": "Filter by content types"
        },
        "filters": {
            "type": "object",
            "description": "Additional search filters"
        },
        "document_id": {
            "type": "string",
            "description": "Document ID (for delete_document or get_statistics)"
        },
        "metadata": {
            "type": "object",
            "description": "Additional metadata for document ingestion"
        }
    },
    "required": ["operation"]
}

POSTGRES_RAG_OUTPUT_SCHEMA = {
    "type": "object",
    "properties": {
        "status": {
            "type": "string",
            "enum": ["success", "error", "duplicate"]
        },
        "result": {
            "type": "object",
            "description": "Operation result data"
        },
        "error": {
            "type": "string",
            "description": "Error message if status is error"
        },
        "processing_time": {
            "type": "number",
            "description": "Time taken to process the operation"
        }
    },
    "required": ["status"]
}